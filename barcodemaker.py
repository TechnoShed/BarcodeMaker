import os
import qrcode
import platform
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWidgets import *
from PyQt5 import uic, QtGui
import pandas as pd

class MyGUI(QMainWindow):

    def __init__(self):
        super(MyGUI, self).__init__()
        uic.loadUi("barcodegui.ui", self)
        self.show()

        self.currentdoc = ""
        self.actionLoad.triggered.connect(self.loadXL)
        self.actionLoad_TXT.triggered.connect(self.loadTXT)
        self.pushButton.clicked.connect(self.savedoc)
        self.pushButton_2.clicked.connect(self.printdoc)

    def loadXL(self):
        options= QFileDialog.Options()
        filename, _ =QFileDialog.getOpenFileName(self ,"Open XL file", "", "Files (*.xlsx)")
        if not filename:
            pass
        else:
            inputdata = pd.read_excel(filename)
            inputdata= inputdata['VIN'].to_string(index=False)
            self.textEdit.setPlainText(inputdata)
  
    def loadTXT(self):
        options= QFileDialog.Options()
        filename, _ =QFileDialog.getOpenFileName(self ,"Open TXT file", "", "Files (*.txt)")
        if not filename:
            pass
        else:
            with open(filename,"r") as f:
                inputtext = f.read()
            self.textEdit.setPlainText(inputtext)
    
    def generatedoc(self, action):
        text2add = self.textEdit.toPlainText().upper().split("\n")
        document= Document()
        section = document.sections[0]
        footer= section.header
        footertext = footer.paragraphs[0]
        footer.text="Barcode Maker (c)2023 K. Winstanley"
        # Loop through the input box , create QRcode, Add to doc file
        for vin in text2add:
            if vin !='':
                print("Creating QRcode for "+ vin)
                qr = qrcode.QRCode(version=1,
                            error_correction=qrcode.constants.ERROR_CORRECT_L,
                            box_size=20,
                            border=2)
                qr.add_data(vin)
                qr.make(fit=True)
                img= qr.make_image(fill_color="Black", back_color="White")
                img.save("currentqr.png")
                document.add_heading(vin,0)
                document.add_paragraph()
                document.add_picture("currentqr.png", width=Inches(1), height=Inches(1))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                document.add_paragraph()
                document.add_picture("currentqr.png", width=Inches(5), height=Inches(5))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                document.add_page_break()
            
        pixmap = QtGui.QPixmap("currentqr.png")
        pixmap = pixmap.scaled(300, 300)
        self.label.setScaledContents(True)
        self.label.setPixmap(pixmap)
        filename = "temp.doc"
        if action =="save":
            options= QFileDialog.Options()
            filename, _ =QFileDialog.getSaveFileName(self ,"Save doc file", "", "Files (*.doc)") 
            if not filename:
                filename = "temp.doc" 
        print("filename "+ filename)
        document.save(filename)
        print("QRcode Creation Process Completed")
    
    def savedoc(self):
        self.generatedoc(action="save")
        
    def printdoc(self):
        self.generatedoc(action="print")
        systype = platform.system()
        
        if systype == "Windows":
            os.startfile("temp.doc", "print")

        if systype == "Linux":
            print("Printing to linux")
    

def main():
    app=QApplication([])
    window = MyGUI()
    app.exec_()

if __name__ == "__main__":
    main()
